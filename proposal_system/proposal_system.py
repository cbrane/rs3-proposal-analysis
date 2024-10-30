# Standard library imports
import ast
import io
import os
import pickle
import random
import tempfile
from datetime import datetime
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

# Third-party imports
import boto3
import pandas as pd
from docx import Document
from PyPDF2 import PdfMerger
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

# Local application imports
import amendment_handler
import email_classifier
import file_classification
import report_generator

# Constants
BUCKET_NAME = "rs3-files"

# 1. S3 Folder and File Check (when a new email is sent, the files will be
#    uploaded to the S3 bucket inside the RS3 folder in the root folder. Once
#    the script has been run, it will 'archive' those files inside a subfolder)
#    The usage of root folder means that for each RS3 number folder, we are
#    only looking to see if there are files in that first folder, not within
#    any subfolders.


def list_folders(bucket_name):
    """
    List all top-level folders in the specified S3 bucket.

    Args:
        bucket_name (str): Name of the S3 bucket to list folders from.

    Returns:
        list: List of folder prefixes (strings) representing top-level folders.
    """
    s3 = boto3.client("s3")
    response = s3.list_objects_v2(
        Bucket=bucket_name,
        Delimiter="/"
    )
    return [
        prefix["Prefix"]
        for prefix in response.get("CommonPrefixes", [])
    ]


def check_notable_files(bucket_name, folder):
    """
    Check for PDF, DOCX, and PKL files in the specified folder.

    Args:
        bucket_name (str): Name of the S3 bucket to check.
        folder (str): Folder prefix to check within the bucket.

    Returns:
        list: List of file keys (strings) for PDF, DOCX and PKL files found.
    """
    s3 = boto3.client("s3")
    response = s3.list_objects_v2(
        Bucket=bucket_name,
        Prefix=folder,
        Delimiter="/"
    )
    return [
        obj["Key"]
        for obj in response.get("Contents", [])
        if (
            obj["Key"] != folder
            and obj["Key"].lower().endswith((".pdf", ".docx", ".pkl"))
        )
    ]


def check_root_files(bucket_name, folder):
    """
    Check for all files in the root of the specified folder.

    Args:
        bucket_name (str): Name of the S3 bucket to check.
        folder (str): Folder prefix to check within the bucket.

    Returns:
        list: List of file keys (strings) found in the root of the folder.
    """
    s3 = boto3.client("s3")
    response = s3.list_objects_v2(
        Bucket=bucket_name,
        Prefix=folder,
        Delimiter="/"
    )
    return [
        obj["Key"]
        for obj in response.get("Contents", [])
        if obj["Key"] != folder
    ]

def process_s3_folders(bucket_name):
    """Process all folders in the S3 bucket and create a DataFrame with file
    information.

    Args:
        bucket_name (str): Name of the S3 bucket to process.

    Returns:
        pd.DataFrame: DataFrame containing file info from the S3 folders.
    """
    s3 = boto3.client("s3")
    folders = list_folders(bucket_name)
    data = []

    for folder in folders:
        root_files = check_root_files(bucket_name, folder)
        notable_files = check_notable_files(bucket_name, folder)
        if not root_files:
            continue

        row = {"Folder": folder, "Files": ", ".join(notable_files)}
        pickle_file = next(
            (f for f in root_files if f.lower().endswith(".pkl")),
            None
        )

        # Add pickle file name as a separate column
        row["Pickle_File"] = pickle_file if pickle_file else ""

        if pickle_file:
            try:
                response = s3.get_object(
                    Bucket=bucket_name,
                    Key=pickle_file
                )
                pickle_content = response["Body"].read()
                pickle_file_obj = io.BytesIO(pickle_content)
                pickle_data = pickle.load(pickle_file_obj)

                if isinstance(pickle_data, dict):
                    row.update(pickle_data)
                else:
                    print(
                        "Warning: Pickle file "
                        f"{pickle_file} does not contain a dictionary"
                    )
            except Exception as e:
                print(
                    f"Error processing pickle file {pickle_file}: "
                    f"{str(e)}"
                )

        data.append(row)

    return pd.DataFrame(data)

def display_rs3_folders_df(rs3_folders_df):
    """Display the DataFrame contents.

    Args:
        rs3_folders_df (pd.DataFrame): DataFrame to display.
    """
    # Set display options
    pd.set_option('display.max_rows', None)
    pd.set_option('display.max_columns', None)
    pd.set_option('display.width', None)
    pd.set_option('display.max_colwidth', None)

    # Print the dataframe
    print(rs3_folders_df.to_string())

    # Reset display options
    pd.reset_option('display.max_rows')
    pd.reset_option('display.max_columns')
    pd.reset_option('display.width')
    pd.reset_option('display.max_colwidth')

def s3_and_df_creation():
    """Create and initialize the global DataFrame with S3 folder info."""
    global rs3_folders_df
    rs3_folders_df = process_s3_folders(BUCKET_NAME)


def print_rs3_folders_df_status():
    """Print the status of the RS3 folders DataFrame."""
    if rs3_folders_df.empty:
        print("No folders with files in the root found.")
        exit()
    print("Files found in the root of folder(s).")


# 2. Run Email Classifier on Folders with files in root folder (each folder
#    will have a .pkl file containing the email that is associated with the
#    files that got dropped into the folder)


def get_pickle_content(bucket_name: str, pickle_file: str) -> dict:
    """Retrieve and load the content of a pickle file from S3.

    Args:
        bucket_name: Name of the S3 bucket containing the pickle file
        pickle_file: Path to the pickle file in the S3 bucket

    Returns:
        Dictionary containing the unpickled data, or None if there was an error
    """
    s3 = boto3.client("s3")
    try:
        response = s3.get_object(Bucket=bucket_name, Key=pickle_file)
        pickle_content = response["Body"].read()
        pickle_file_obj = io.BytesIO(pickle_content)
        return pickle.load(pickle_file_obj)
    except Exception as e:
        print(f"Error processing pickle file {pickle_file}: {str(e)}")
        return None


def run_email_classifier():
    """Run email classifier on each row of the DataFrame."""
    for index, row in rs3_folders_df.iterrows():
        pickle_file = row["Pickle_File"]
        if pickle_file:
            pickle_data = get_pickle_content("rs3-files", pickle_file)

            if pickle_data:
                # Combine subject and body for email classification
                email_content = (
                    f"Subject: {pickle_data.get('subject', '')}\n\n"
                    f"Body:\n{pickle_data.get('body', '')}"
                )

                # Run email classifier
                is_new_report = email_classifier.classify_email(email_content)

                # Update the DataFrame with the classification result
                rs3_folders_df.at[index, "is_new_report"] = is_new_report
                rs3_folders_df.at[index, "Email Classifier Result"] = (
                    "Report Needed" if is_new_report else "Amendment/Other"
                )
            else:
                rs3_folders_df.at[index, "is_new_report"] = None
                rs3_folders_df.at[index, "Email Classifier Result"] = (
                    "Error processing pickle file"
                )


def print_email_classifier_results():
    """Print the email classifier results for each folder."""
    for index, row in rs3_folders_df.iterrows():
        if pd.notna(row.get("is_new_report")) and row["is_new_report"]:
            print(f"Folder: {row['Folder']}")
            print(f"Email Classifier Result: {row['Email Classifier Result']}")


# 3a. 'Report Needed' Pipeline - (If new_report = True (report creation needed))
#     Run File Classifier for Folders/Emails that have been classified as a new
#     report


def classify_files_in_folder(
    bucket_name: str,
    folder: str,
    files: str,
) -> dict:
    """Classify files in the folder as RS3 reports or not.

    Args:
        bucket_name: Name of the S3 bucket
        folder: Folder path in the bucket
        files: Comma-separated list of files

    Returns:
        Dictionary of classified files with their status
    """
    classified_files = {}
    for file in files.split(", "):
        # Only process .docx and .pdf files
        if file.lower().endswith((".docx", ".pdf")):
            is_rs3, reason = file_classification.is_rs3_report_by_name(file)
            classified_files[file] = {"is_rs3": is_rs3, "reason": reason}
    return classified_files


def run_file_classifier():
    """Run file classifier on each row of the DataFrame."""
    for index, row in rs3_folders_df.iterrows():
        if pd.notna(row.get("is_new_report")) and row["is_new_report"]:
            folder = row["Folder"]
            files = row["Files"]

            # Classify files in the folder
            classified_files = classify_files_in_folder(
                "rs3-files",
                folder,
                files
            )

            # Add the classification results to the DataFrame
            rs3_folders_df.at[index, "Classified Files"] = str(classified_files)

            # Count RS3 files
            rs3_file_count = sum(
                1 for file_info in classified_files.values()
                if file_info["is_rs3"]
            )
            rs3_folders_df.at[index, "RS3 File Count"] = rs3_file_count


def get_rs3_files(row):
    """Get the list of RS3 files from the classified files in a row.

    Args:
        row (pandas.Series): Row from the DataFrame

    Returns:
        list: List of RS3 files
    """
    classified_files = ast.literal_eval(row["Classified Files"])
    return [file for file, info in classified_files.items() if info["is_rs3"]]


def print_rs3_files():
    """Iterate through folders and display any RS3 files that were found.
    
    Prints the folder path and lists all RS3 files within each folder that
    requires a new report. If no RS3 files are found in a folder, prints
    a message indicating this.
    """
    for index, row in rs3_folders_df.iterrows():
        if row["is_new_report"]:
            print(f"\nFolder: {row['Folder']}")
            rs3_files = get_rs3_files(row)
            if rs3_files:
                print(f"RS3 files found ({len(rs3_files)}):")
                for file in rs3_files:
                    print(f"- {file}")
            else:
                print("No RS3 files found.")

# 4a. 'Report Needed' Pipeline, (If new_report = True (report creation needed))
# After the File Classifier has been run, feed the file (or combine multiple files)
# into report_generator.py to create a new RS3 report.


def download_file(bucket_name: str, file_key: str) -> io.BytesIO:
    """Download a file from S3 and return its content as BytesIO object.

    Args:
        bucket_name: Name of the S3 bucket to download from
        file_key: Key/path of the file in the S3 bucket

    Returns:
        BytesIO object containing the downloaded file content
    """
    s3 = boto3.client("s3")
    response = s3.get_object(Bucket=bucket_name, Key=file_key)
    return io.BytesIO(response["Body"].read())


def combine_docx(files: list) -> Document:
    """Combine multiple DOCX files into a single Document object.

    Args:
        files: List of file objects to combine

    Returns:
        Document: Combined document containing content from all input files
    """
    combined = Document()
    for file in files:
        doc = Document(file)
        for element in doc.element.body:
            combined.element.body.append(element)
    return combined


def combine_pdf(files: list) -> io.BytesIO:
    """Combine multiple PDF files into a single PDF.

    Args:
        files: List of PDF file objects to combine

    Returns:
        io.BytesIO: Combined PDF file as a BytesIO object
    """
    merger = PdfMerger()
    for file in files:
        merger.append(file)
    output = io.BytesIO()
    merger.write(output)
    output.seek(0)
    return output


def docx_to_pdf(docx_content: io.BytesIO) -> io.BytesIO:
    """Convert a DOCX file to PDF format.

    Args:
        docx_content: BytesIO object containing the DOCX file

    Returns:
        io.BytesIO: Converted PDF file as a BytesIO object
    """
    doc = Document(docx_content)
    pdf_buffer = io.BytesIO()
    pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = []

    for para in doc.paragraphs:
        flowables.append(Paragraph(para.text, styles["Normal"]))

    pdf.build(flowables)
    pdf_buffer.seek(0)
    return pdf_buffer


def combine_files(bucket_name: str, rs3_files: list) -> tuple:
    """Combine multiple files (DOCX or PDF) into a single PDF.

    Args:
        bucket_name (str): Name of the S3 bucket
        rs3_files (list): List of files to combine

    Returns:
        tuple: Combined file as BytesIO object and its file type
    """
    file_contents = [download_file(bucket_name, file) for file in rs3_files]

    if all(file.lower().endswith(".docx") for file in rs3_files):
        combined = combine_docx(file_contents)
        output = io.BytesIO()
        combined.save(output)
        output.seek(0)
        return docx_to_pdf(output), "pdf"
    elif all(file.lower().endswith(".pdf") for file in rs3_files):
        return combine_pdf(file_contents), "pdf"
    else:
        # If mixed file types, convert all to PDF and then combine
        pdf_files = []
        for file, content in zip(rs3_files, file_contents):
            if file.lower().endswith(".docx"):
                pdf_files.append(docx_to_pdf(content))
            else:
                pdf_files.append(content)
        return combine_pdf(pdf_files), "pdf"


def save_bytesio_to_temp_file(bytesio, file_type):
    """Save a BytesIO object to a temporary file.

    Args:
        bytesio (BytesIO): BytesIO object to save
        file_type (str): File extension for the temporary file

    Returns:
        str: Path to the temporary file
    """
    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=f".{file_type}"
    )
    temp_file.write(bytesio.getvalue())
    temp_file.close()
    return temp_file.name


def download_and_save_temp_file(bucket_name: str, file_key: str) -> str:
    """Download a file from S3 and save it as a temporary file.

    Args:
        bucket_name: Name of the S3 bucket
        file_key: Key of the file in the bucket

    Returns:
        str: Path to the temporary file
    """
    file_content = download_file(bucket_name, file_key)
    file_type = file_key.split(".")[-1]
    return save_bytesio_to_temp_file(file_content, file_type)


# Prepare files for report generation
def prepare_files_for_report_generation():
    """Prepare files for report generation."""
    for index, row in rs3_folders_df.iterrows():
        if row["Email Classifier Result"] == "Report Needed":
            rs3_files = get_rs3_files(row)
            pickle_file = row["Pickle_File"]

            if rs3_files:
                if len(rs3_files) > 1:
                    combined_file, file_type = combine_files(
                        "rs3-files",
                        rs3_files
                    )
                    rs3_temp_file_path = save_bytesio_to_temp_file(
                        combined_file,
                        file_type
                    )
                else:
                    file_content = download_file("rs3-files", rs3_files[0])
                    if rs3_files[0].lower().endswith(".docx"):
                        pdf_content = docx_to_pdf(file_content)
                        rs3_temp_file_path = save_bytesio_to_temp_file(
                            pdf_content,
                            "pdf"
                        )
                    else:
                        rs3_temp_file_path = save_bytesio_to_temp_file(
                            file_content,
                            "pdf"
                        )

                # Download and save pickle file
                pickle_temp_file_path = download_and_save_temp_file(
                    "rs3-files",
                    pickle_file
                )

                rs3_folders_df.at[index, "RS3 Temp File Path"] = rs3_temp_file_path
                rs3_folders_df.at[index, "Pickle Temp File Path"] = pickle_temp_file_path
            else:
                print(f"No RS3 files found for folder: {row['Folder']}")
                rs3_folders_df.at[index, "RS3 Temp File Path"] = None
                rs3_folders_df.at[index, "Pickle Temp File Path"] = None


def generate_reports():
    """Generate reports for rows classified as 'Report Needed'."""
    for index, row in rs3_folders_df.iterrows():
        if (
            row["Email Classifier Result"] == "Report Needed"
            and row["RS3 Temp File Path"]
            and row["Pickle Temp File Path"]
        ):
            try:
                # Generate report
                rs3_file, result = report_generator.generate_report(
                    row["RS3 Temp File Path"],
                    row["Pickle Temp File Path"]
                )

                # Store the result in the DataFrame
                rs3_folders_df.at[index, "Report Generation File"] = rs3_file
                rs3_folders_df.at[index, "Recommend Bid"] = result

            except Exception as e:
                print(f"Error generating report for row {index}: {str(e)}")
                rs3_folders_df.at[index, "Report Generation File"] = f"Error: {str(e)}"
                rs3_folders_df.at[index, "Recommend Bid"] = f"Error: {str(e)}"


# 3b. 'Amendment/Other' Pipeline - (If new_report = False)
# Run Amendment Handler for emails/folders classified as not needing reports


def download_file(bucket_name: str, file_key: str) -> io.BytesIO:
    """Download a file from S3 and return its content as BytesIO object.

    Args:
        bucket_name: Name of the S3 bucket
        file_key: Key of the file in S3

    Returns:
        BytesIO object containing the file content
    """
    s3 = boto3.client("s3")
    response = s3.get_object(Bucket=bucket_name, Key=file_key)
    return io.BytesIO(response["Body"].read())


def save_bytesio_to_temp_file(bytesio: io.BytesIO, file_type: str) -> str:
    """Save a BytesIO object to a temporary file and return the file path.

    Args:
        bytesio: BytesIO object containing file content
        file_type: Extension of the file (e.g. 'pdf', 'docx')

    Returns:
        Path to the saved temporary file
    """
    temp_file = tempfile.NamedTemporaryFile(
        delete=False, 
        suffix=f".{file_type}"
    )
    temp_file.write(bytesio.getvalue())
    temp_file.close()
    return temp_file.name

def download_and_save_temp_file(bucket_name: str, file_key: str) -> str:
    """Download a file from S3 and save it as a temporary file.

    Args:
        bucket_name: Name of the S3 bucket
        file_key: Key of the file in S3

    Returns:
        str: Path to the saved temporary file
    """
    file_content = download_file(bucket_name, file_key)
    file_type = file_key.split(".")[-1]
    return save_bytesio_to_temp_file(file_content, file_type)


# Prepare files for amendment handling
def prepare_files_for_amendment_handling():
    """Prepare files for amendment handling."""
    for index, row in rs3_folders_df.iterrows():
        if row["Email Classifier Result"] == "Amendment/Other":
            pickle_file = row["Pickle_File"]

            # Download and save pickle file
            pickle_temp_file_path = download_and_save_temp_file(
                "rs3-files",
                pickle_file
            )

            rs3_folders_df.at[index, "Pickle Temp File Path"] = pickle_temp_file_path


def generate_amendment_reports():
    """Generate reports for amendments."""
    for index, row in rs3_folders_df.iterrows():
        if (row["Email Classifier Result"] == "Amendment/Other" and
                row["Pickle Temp File Path"]):
            try:
                # Generate report
                handler = amendment_handler.AmendmentHandler(
                    row["Pickle Temp File Path"])
                file_name = handler.process_amendment()

                # Store the result in the DataFrame
                rs3_folders_df.at[index, "Report Generation File"] = file_name

            except Exception as e:
                print(f"Error generating report for row {index}: {str(e)}")
                rs3_folders_df.at[index, "Report Generation File"] = f"Error: {str(e)}"


# Save Report files into S3 bucket
# For each row in the DataFrame, save the report file in 'Report Generation
# File' to the row in the column 'Folder'


def save_reports_to_s3(df, bucket_name):
    """Save generated report files to S3 bucket."""
    s3 = boto3.client("s3")

    for index, row in df.iterrows():
        report_file = row.get("Report Generation File")
        folder = row.get("Folder")

        if report_file and folder and not report_file.startswith("Error:"):
            local_file_path = report_file
            file_name = os.path.basename(local_file_path)
            s3_key = f"{folder}{file_name}"

            try:
                with open(local_file_path, "rb") as file:
                    s3.upload_fileobj(file, bucket_name, s3_key)
                print(f"Successfully uploaded {file_name} to {s3_key}")
            except Exception as e:
                print(f"Error uploading {file_name} to S3: {str(e)}")


def send_email(df):
    """Send email with SES using generated report as attachment."""
    region_name = os.getenv("AWS_REGION", "us-east-1")
    client = boto3.client("ses", region_name=region_name)
    sender_email = "connor@connorraney.com"
    recipient_email = "connor@connorraney.com"

    for _, row in df.iterrows():
        if row.get("is_new_report", False):
            msg = MIMEMultipart()
            rs3_number = row.get("Folder", "Unknown RS3").rstrip("/")

            msg["Subject"] = f"RS3 Report: {rs3_number}"
            msg["From"] = sender_email
            msg["To"] = recipient_email

            body = f"Here is the report for {rs3_number} attached."
            msg.attach(MIMEText(body, "plain"))

            attachment_path = row.get("Report Generation File")
            if not attachment_path or not os.path.exists(attachment_path):
                print(
                    f"Attachment not found for {rs3_number}: {attachment_path}"
                )
                continue

            with open(attachment_path, "rb") as attachment:
                part = MIMEApplication(attachment.read())
                part.add_header(
                    "Content-Disposition",
                    "attachment",
                    filename=os.path.basename(attachment_path),
                )
                msg.attach(part)

            try:
                response = client.send_raw_email(
                    Source=sender_email,
                    Destinations=[recipient_email],
                    RawMessage={"Data": msg.as_string()},
                )
                print(
                    f"Email sent successfully for {rs3_number}. "
                    f"Message ID: {response['MessageId']}"
                )
            except Exception as e:
                print(f"Failed to send email for {rs3_number}: {str(e)}")
                break


# Archive Files - Go through all RS3 folders in the "rs3-files" S3 bucket & archive


def archive_root_files(bucket_name: str) -> None:
    """Archive root files in all folders of the specified S3 bucket.

    Args:
        bucket_name: Name of the S3 bucket
    """
    s3 = boto3.client("s3")

    try:
        folders = list_folders(bucket_name)
    except Exception as e:
        print(f"Error listing folders in bucket {bucket_name}: {str(e)}")
        return

    for folder in folders:
        print(f"Processing folder: {folder}")
        try:
            response = s3.list_objects_v2(
                Bucket=bucket_name,
                Prefix=folder,
                Delimiter="/"
            )

            root_files = [
                obj["Key"]
                for obj in response.get("Contents", [])
                if obj["Key"] != folder
                and "/" not in obj["Key"].replace(folder, "", 1)
            ]

            if root_files:
                now = datetime.now()
                archive_folder = now.strftime("%m-%d-%Y-%H%M-archive")

                for file in root_files:
                    try:
                        old_key = file
                        new_key = (
                            f"{folder}{archive_folder}/"
                            f"{file.split('/')[-1]}"
                        )
                        s3.copy_object(
                            Bucket=bucket_name,
                            CopySource=f"{bucket_name}/{old_key}",
                            Key=new_key,
                        )
                        s3.delete_object(
                            Bucket=bucket_name,
                            Key=old_key
                        )
                        print(f"Archived file: {old_key} to {new_key}")
                    except Exception as e:
                        print(f"Error archiving file {old_key}: {str(e)}")
            else:
                print(f"No root files found in {folder}")

        except Exception as e:
            print(f"Error processing folder {folder}: {str(e)}")

    print("Archiving complete.")


def unarchive_files(num_folders):
    """Unarchive files from random folders in the S3 bucket.

    Selects random folders and moves non-report files from archive
    locations to the root of their respective folders. Includes safety
    checks before deleting originals.

    Args:
        num_folders (int): Number of random folders to unarchive from.

    Returns:
        None: Prints progress and summary information.
    """
    s3 = boto3.client("s3")

    all_folders = list_folders(BUCKET_NAME)
    num_folders = min(num_folders, len(all_folders))
    selected_folders = random.sample(all_folders, num_folders)

    total_unarchived = 0

    for random_folder in selected_folders:
        # List contents of the random folder to find the archive subfolder
        response = s3.list_objects_v2(
            Bucket=BUCKET_NAME,
            Prefix=random_folder,
            Delimiter="/"
        )
        subfolders = [
            content["Prefix"]
            for content in response.get("CommonPrefixes", [])
        ]

        # Find the archive subfolder (assuming it ends with '-archive')
        archive_subfolder = next(
            (sf for sf in subfolders if sf.endswith("-archive/")),
            None
        )

        if archive_subfolder:
            # List files in the archive subfolder
            archive_files = check_root_files(BUCKET_NAME, archive_subfolder)

            for file in archive_files:
                if "-report" not in file:  # Skip report files
                    old_key = file
                    new_key = f"{random_folder}{file.split('/')[-1]}"

                    try:
                        # Copy the file to the new location
                        s3.copy_object(
                            Bucket=BUCKET_NAME,
                            CopySource=f"{BUCKET_NAME}/{old_key}",
                            Key=new_key,
                        )

                        # Check if the file was successfully copied
                        s3.head_object(Bucket=BUCKET_NAME, Key=new_key)

                        # If the check passes, delete the original file
                        s3.delete_object(Bucket=BUCKET_NAME, Key=old_key)

                        total_unarchived += 1
                        print(f"Successfully unarchived: {old_key} to {new_key}")

                    except Exception as e:
                        print(f"Error unarchiving {old_key}: {str(e)}")

            print(f"Unarchived files from folder {random_folder}")
        else:
            print(f"No archive subfolder found in {random_folder}")

    print(
        f"Unarchiving complete. Total {total_unarchived} files unarchived "
        f"from {num_folders} folders."
    )


def manage_folder_archives(bucket_name, folder, action="archive"):
    """Archive or unarchive files in a specific folder in the S3 bucket.

    Args:
        bucket_name (str): Name of the S3 bucket.
        folder (str): Specific folder to archive or unarchive.
        action (str): 'archive' to move files to archive,
                     'unarchive' to restore files.
    """
    s3 = boto3.client("s3")

    print(f"Processing folder: {folder}")
    try:
        if action == "archive":
            response = s3.list_objects_v2(
                Bucket=bucket_name,
                Prefix=folder,
                Delimiter="/"
            )

            files = [
                obj["Key"]
                for obj in response.get("Contents", [])
                if obj["Key"] != folder
            ]
            subfolders = [
                content["Prefix"]
                for content in response.get("CommonPrefixes", [])
            ]

            if files or subfolders:
                now = datetime.now()
                archive_folder = now.strftime("%m-%d-%Y-%H%M-archive")

                for file in files:
                    if not file.endswith("/"):
                        try:
                            old_key = file
                            new_key = (
                                f"{folder}{archive_folder}/"
                                f"{file.split('/')[-1]}"
                            )
                            s3.copy_object(
                                Bucket=bucket_name,
                                CopySource=f"{bucket_name}/{old_key}",
                                Key=new_key,
                            )
                            s3.delete_object(Bucket=bucket_name, Key=old_key)
                            print(f"Archived file: {old_key} to {new_key}")
                        except Exception as e:
                            print(f"Error archiving file {old_key}: {str(e)}")

                for subfolder in subfolders:
                    try:
                        subfolder_name = subfolder.split("/")[-2]
                        subfolder_objects = s3.list_objects_v2(
                            Bucket=bucket_name,
                            Prefix=subfolder
                        )

                        for obj in subfolder_objects.get("Contents", []):
                            try:
                                old_key = obj["Key"]
                                new_key = (
                                    f"{folder}{archive_folder}/"
                                    f"{subfolder_name}/"
                                    f"{old_key.split('/')[-1]}"
                                )
                                s3.copy_object(
                                    Bucket=bucket_name,
                                    CopySource=f"{bucket_name}/{old_key}",
                                    Key=new_key,
                                )
                                s3.delete_object(
                                    Bucket=bucket_name,
                                    Key=old_key
                                )
                                print(
                                    f"Archived object from subfolder: "
                                    f"{old_key} to {new_key}"
                                )
                            except Exception as e:
                                print(
                                    f"Error archiving object {old_key} "
                                    f"from subfolder: {str(e)}"
                                )
                    except Exception as e:
                        print(
                            f"Error processing subfolder {subfolder}: "
                            f"{str(e)}"
                        )
            else:
                print(f"No files or subfolders found in {folder}")

            print(f"Archiving complete for folder {folder}")

        elif action == "unarchive":
            response = s3.list_objects_v2(
                Bucket=bucket_name,
                Prefix=folder
            )

            for obj in response.get("Contents", []):
                if obj["Key"] != folder and not obj["Key"].endswith("/"):
                    old_key = obj["Key"]
                    file_name = old_key.split("/")[-1]

                    # Skip files with '-report' in the name
                    if "-report" in file_name:
                        continue

                    # Move file one folder up
                    new_key = "/".join(old_key.split("/")[:-2] + [file_name])

                    try:
                        s3.copy_object(
                            Bucket=bucket_name,
                            CopySource=f"{bucket_name}/{old_key}",
                            Key=new_key,
                        )
                        s3.delete_object(Bucket=bucket_name, Key=old_key)
                        print(f"Unarchived file: {old_key} to {new_key}")
                    except Exception as e:
                        print(f"Error unarchiving file {old_key}: {str(e)}")

            print(f"Unarchiving complete for folder {folder}")

        else:
            print(f"Invalid action: {action}. Use 'archive' or 'unarchive'.")

    except Exception as e:
        print(f"Error processing folder {folder}: {str(e)}")